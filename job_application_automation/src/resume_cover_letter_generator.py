"""
Resume and cover letter generator module using Llama 4 Mevrick.
This module provides functionality to generate personalized resumes and cover letters
based on job descriptions and user profiles.
"""

import os
import json
import logging
import asyncio
import requests
import time
from typing import Dict, Any, Optional, Union, List, Tuple
from pathlib import Path
from datetime import datetime
from enum import Enum

# Document generation libraries
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches

# AI/LLM integration
from llama_cpp import Llama  # type: ignore

# Import configuration
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config.llama_config import LlamaConfig

# Set up logging with absolute path for the log file
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
log_file_path = os.path.join(project_root, "data", "resume_cover_letter_generator.log")

# Ensure log directory exists
os.makedirs(os.path.dirname(log_file_path), exist_ok=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file_path),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Enum for cover letter templates
class CoverLetterTemplate(str, Enum):
    STANDARD = "standard"
    TECHNICAL = "technical"
    CREATIVE = "creative"
    EXECUTIVE = "executive"
    CAREER_CHANGE = "career_change"
    REFERRAL = "referral"

# Template manager for cover letters
class CoverLetterTemplateManager:
    """
    Manages different cover letter templates for various scenarios.
    """
    
    def __init__(self):
        """Initialize the template manager"""
        self.templates_dir = Path("../templates")
        self.template_paths = {
            CoverLetterTemplate.STANDARD: self.templates_dir / "cover_letter_template.docx",
            CoverLetterTemplate.CREATIVE: self.templates_dir / "cover_letter_creative_template.docx",
            CoverLetterTemplate.TECHNICAL: self.templates_dir / "cover_letter_technical_template.docx",
            CoverLetterTemplate.EXECUTIVE: self.templates_dir / "cover_letter_executive_template.docx",
            CoverLetterTemplate.CAREER_CHANGE: self.templates_dir / "cover_letter_career_change_template.docx",
            CoverLetterTemplate.REFERRAL: self.templates_dir / "cover_letter_referral_template.docx"
        }
        
        # Template prompts for different cover letter styles
        self.template_prompts = {
            CoverLetterTemplate.STANDARD: """
            Based on the following job description and candidate resume, create a personalized cover letter.
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Company Information:
            {company_info}
            
            Create a compelling cover letter that:
            1. Addresses the hiring manager professionally (by name if available, otherwise "Hiring Manager")
            2. Expresses enthusiasm for the position and company
            3. Highlights 2-3 most relevant experiences/skills from the resume that match the job requirements
            4. Explains why the candidate is a good fit for the company culture
            5. Includes a strong closing paragraph with a call to action
            
            The tone should be professional yet personable, confident but not arrogant.
            Format the letter properly with current date, recipient's information, and proper salutation and closing.
            """,
            
            CoverLetterTemplate.CREATIVE: """
            Based on the following job description and candidate resume, create a creative and engaging cover letter.
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Company Information:
            {company_info}
            
            Create an attention-grabbing cover letter that:
            1. Opens with a unique hook or story that relates to the position
            2. Conveys passion for the company's mission or industry
            3. Demonstrates creativity and innovation through specific examples
            4. Shows personality while maintaining professionalism
            5. Closes with enthusiasm and a clear call to action
            
            The tone should be engaging, conversational, and show genuine excitement.
            This template works well for creative industries (design, marketing, content creation).
            """,
            
            CoverLetterTemplate.TECHNICAL: """
            Based on the following job description and candidate resume, create a technically-focused cover letter.
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Company Information:
            {company_info}
            
            Create a precise, technically-oriented cover letter that:
            1. Addresses the technical requirements of the role directly
            2. Highlights specific technical skills, certifications, and achievements
            3. Demonstrates problem-solving abilities through concrete examples
            4. Shows understanding of the company's technical challenges
            5. Maintains a professional, detail-oriented tone throughout
            
            Focus on quantifiable achievements and specific technical contributions.
            Include relevant technical keywords from the job description.
            This template works well for engineering, development, and IT positions.
            """,
            
            CoverLetterTemplate.EXECUTIVE: """
            Based on the following job description and candidate resume, create an executive-level cover letter.
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Company Information:
            {company_info}
            
            Create a powerful executive cover letter that:
            1. Opens with a strong statement of value proposition
            2. Highlights strategic leadership and vision
            3. Demonstrates business impact through measurable achievements
            4. Shows industry expertise and key leadership qualities
            5. Conveys confidence and executive presence
            
            The tone should be authoritative, strategic, and refined.
            Focus on high-level achievements and leadership capabilities.
            This template works well for C-level, director, and senior management positions.
            """,
            
            CoverLetterTemplate.CAREER_CHANGE: """
            Based on the following job description and candidate resume, create a cover letter for someone changing careers.
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Company Information:
            {company_info}
            
            Create a persuasive cover letter that:
            1. Acknowledges the career transition directly but confidently
            2. Emphasizes transferable skills that apply to the new role
            3. Explains the motivation for the career change
            4. Connects past achievements to potential future contributions
            5. Shows enthusiasm and commitment to the new field
            
            The tone should be confident and forward-looking.
            Focus on how the candidate's diverse experience offers unique value.
            This template works well for career-changers and those with non-traditional backgrounds.
            """,
            
            CoverLetterTemplate.REFERRAL: """
            Based on the following job description and candidate resume, create a cover letter that mentions a referral.
            
            Job Description:
            {job_description}
            
            Candidate Resume:
            {candidate_resume}
            
            Company Information:
            {company_info}
            
            Referral Information:
            {referral_info}
            
            Create an effective referral-based cover letter that:
            1. Mentions the referrer in the opening paragraph
            2. Explains the connection to the referrer and why they recommended this role
            3. Builds on the credibility established by the referral
            4. Highlights relevant skills and experiences that match the job requirements
            5. Ends with appreciation and a clear call to action
            
            The tone should be warm yet professional.
            Leverage the referral appropriately without over-relying on it.
            This template works well when someone inside the company has referred the candidate.
            """
        }
        
    def get_template_path(self, template_type: CoverLetterTemplate) -> Path:
        """
        Get the path to a specific template file.
        
        Args:
            template_type: The type of cover letter template to use.
            
        Returns:
            Path object to the template file.
        """
        # Use the standard template if the requested one doesn't exist
        if not self.template_paths[template_type].exists():
            logger.warning(f"Template {template_type.value} not found. Using standard template.")
            return self.template_paths[CoverLetterTemplate.STANDARD]
        
        return self.template_paths[template_type]
    
    def get_template_prompt(self, template_type: CoverLetterTemplate) -> str:
        """
        Get the prompt for a specific template type.
        
        Args:
            template_type: The type of cover letter template to use.
            
        Returns:
            String prompt for the selected template.
        """
        return self.template_prompts.get(template_type, self.template_prompts[CoverLetterTemplate.STANDARD])
    
    def select_best_template(self, job_description: str, candidate_profile: Dict[str, Any]) -> CoverLetterTemplate:
        """
        Automatically select the most appropriate template based on the job and candidate.
        
        Args:
            job_description: The job description text.
            candidate_profile: The candidate's profile information.
            
        Returns:
            The most appropriate CoverLetterTemplate type.
        """
        job_desc_lower = job_description.lower()
        
        # Check for executive positions
        if any(term in job_desc_lower for term in ['ceo', 'cto', 'cfo', 'chief', 'director', 'vp', 'vice president', 'head of']):
            return CoverLetterTemplate.EXECUTIVE
            
        # Check for technical positions
        if any(term in job_desc_lower for term in ['engineer', 'developer', 'programmer', 'architect', 'data scientist']):
            return CoverLetterTemplate.TECHNICAL
            
        # Check for creative positions
        if any(term in job_desc_lower for term in ['design', 'creative', 'writer', 'artist', 'content', 'marketing']):
            return CoverLetterTemplate.CREATIVE
            
        # Default to standard
        return CoverLetterTemplate.STANDARD


class ResumeGenerator:
    """
    Class for generating personalized resumes and cover letters using LLM.
    """

    def __init__(self, config: Optional[LlamaConfig] = None):
        """
        Initialize the ResumeGenerator with configuration settings.
        
        Args:
            config: Configuration settings for LLM integration.
                   If None, default settings will be used.
        """
        self.config = config or LlamaConfig()
        self._setup_llm()
        self.template_manager = CoverLetterTemplateManager()
        self.llm_available = False
        self.api_available = False
        
    def _setup_llm(self) -> None:
        """Set up the LLM for resume and cover letter generation."""
        try:
            # Create models directory if it doesn't exist 
            models_dir = os.path.join(project_root, "models")
            os.makedirs(models_dir, exist_ok=True)

            # Create data directories if they don't exist
            output_dir = os.path.join(project_root, "data", "generated_cover_letters")
            os.makedirs(output_dir, exist_ok=True)
            
            # Check if model file exists
            model_path = os.path.join(project_root, "models", "llama-4-mevrick")
            if not os.path.exists(model_path):
                logger.info("Model file not found at %s - will use template-based generation instead", model_path)
                self.llm_available = False
                return
                
            try:
                # Try to load the model
                import llama_cpp
                
                self.llm = llama_cpp.Llama(
                    model_path=model_path,
                    n_ctx=4096,
                    n_threads=self.config.n_threads,
                    n_gpu_layers=self.config.n_gpu_layers
                )
                self.llm_available = True
                logger.info("LLM initialized successfully")
                
            except ImportError:
                logger.warning("llama_cpp package not installed - falling back to template-based generation")
                self.llm_available = False
                
        except Exception as e:
            logger.warning("Error setting up LLM: %s - will use template-based generation instead", e)
            self.llm_available = False
    
    def _test_api_connection(self, api_config: Dict[str, Any]) -> bool:
        """Test connection to the API provider."""
        try:
            api_base = api_config.get("api_base", "")
            api_key = api_config.get("api_key", "")
            model = api_config.get("model", "")
            timeout = api_config.get("timeout", 10)
            
            if not api_base or not api_key or not model:
                logger.error("Missing API configuration parameters")
                return False
            
            # Prepare a minimal test request
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            
            # Different request format based on provider
            if "openrouter" in api_base:
                # OpenRouter endpoint test
                endpoint = f"{api_base.rstrip('/')}/models"
                response = requests.get(
                    endpoint,
                    headers=headers,
                    timeout=timeout
                )
            else:
                # Groq or OpenAI-compatible endpoint test
                endpoint = f"{api_base.rstrip('/')}/models"
                response = requests.get(
                    endpoint,
                    headers=headers,
                    timeout=timeout
                )
            
            if response.status_code == 200:
                return True
            else:
                logger.error(f"API test failed with status code: {response.status_code}, message: {response.text}")
                return False
                
        except Exception as e:
            logger.error(f"API connection test failed: {e}")
            return False
    
    def _generate_text_with_api(self, prompt: str) -> str:
        """Generate text using the API."""
        api_config = self.config.get_api_config()
        
        if not api_config or not self.api_available:
            logger.error("API not available for text generation")
            return "API error: Could not generate text."
        
        try:
            api_base = api_config.get("api_base", "")
            api_key = api_config.get("api_key", "")
            model = api_config.get("model", "")
            timeout = api_config.get("timeout", 60)
            
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
            
            # Add http-referer header for OpenRouter
            if "openrouter" in api_base:
                headers["HTTP-Referer"] = "https://github.com/job-application-automation"
                headers["X-Title"] = "Job Application Automation Tool"
            
            # Prepare the request payload
            payload = {
                "model": model,
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "temperature": self.config.temperature,
                "top_p": self.config.top_p,
                "max_tokens": min(4000, self.config.context_length)
            }
            
            # Send the request
            endpoint = f"{api_base.rstrip('/')}/chat/completions"
            logger.info(f"Sending request to {endpoint} with model {model}")
            
            response = requests.post(
                endpoint,
                headers=headers,
                json=payload,
                timeout=timeout
            )
            
            # Process the response
            if response.status_code == 200:
                response_json = response.json()
                if self.config.api_provider == "openrouter":
                    generated_text = response_json.get("choices", [{}])[0].get("message", {}).get("content", "")
                    # Log usage information if available
                    if "usage" in response_json:
                        usage = response_json["usage"]
                        logger.info(f"API usage: prompt_tokens={usage.get('prompt_tokens')}, completion_tokens={usage.get('completion_tokens')}")
                else:
                    # Groq or other OpenAI-compatible APIs
                    generated_text = response_json.get("choices", [{}])[0].get("message", {}).get("content", "")
                    # Log usage information if available
                    if "usage" in response_json:
                        usage = response_json["usage"]
                        logger.info(f"API usage: prompt_tokens={usage.get('prompt_tokens')}, completion_tokens={usage.get('completion_tokens')}")
                
                return generated_text
            else:
                logger.error(f"API request failed with status code: {response.status_code}, message: {response.text}")
                return f"API error: {response.status_code} - Could not generate text."
                
        except Exception as e:
            logger.error(f"Error generating text with API: {e}")
            return f"API error: {str(e)}"
            
    def _generate_text_with_llm(self, prompt: str) -> str:
        """Generate text using the local LLM."""
        if not self.llm_available:
            logger.warning("LLM not available for text generation")
            # Use template-based fallback
            return self._template_based_fallback(prompt)
            
        try:
            # Process with local Llama model
            if self.llm:
                output = self.llm(
                    prompt,
                    max_tokens=self.config.context_length,
                    temperature=self.config.temperature,
                    top_p=self.config.top_p,
                    echo=False
                )
                generated_text = output.get("choices", [{}])[0].get("text", "").strip()
                return generated_text
            else:
                return self._template_based_fallback(prompt)
                
        except Exception as e:
            logger.error(f"Error generating with local LLM: {e}")
            return self._template_based_fallback(prompt)
            
    def _template_based_fallback(self, prompt: str) -> str:
        """Generate text based on templates when LLM is not available."""
        # Simple template-based generation logic here
        logger.warning("Using template-based fallback for text generation")
        
        # Identify the type of content needed based on prompt keywords
        if "resume" in prompt.lower() or "skills section" in prompt.lower():
            return "Skills:\n- Technical skills relevant to the position\n- Soft skills like communication and teamwork\n- Industry-specific knowledge\n- Tools and methodologies"
        elif "cover letter" in prompt.lower():
            return "Dear Hiring Manager,\n\nI am writing to express my interest in the position. With my background and experience, I believe I would be a valuable addition to your team.\n\nThank you for your consideration.\n\nSincerely,\n[Your Name]"
        else:
            return "Generated content based on the provided information."
                
    def generate_text(self, prompt: str) -> str:
        """Generate text using available methods (API or local LLM)."""
        # Choose the appropriate generation method
        if self.config.use_api and self.api_available:
            logger.info(f"Generating text with {self.config.api_provider} API")
            return self._generate_text_with_api(prompt)
        elif self.llm_available:
            logger.info("Generating text with local LLM")
            return self._generate_text_with_llm(prompt)
        else:
            logger.warning("No text generation method available, using fallback")
            return self._template_based_fallback(prompt)

    # The remaining methods can use self.generate_text instead of directly using self.llm
    
    def generate_resume(self, 
                  job_description: str, 
                  candidate_profile: Dict[str, Any],
                  output_path: Optional[str] = None) -> Tuple[str, str]:
        """
        Generate a personalized resume based on a job description and candidate profile.
        
        Args:
            job_description: The job description to tailor the resume to.
            candidate_profile: The candidate's profile data.
            output_path: Optional path to save the generated resume.
                         If None, a default path will be used.
                         
        Returns:
            A tuple containing (file_path, resume_content).
        """
        try:
            prompt = self._prepare_resume_prompt(job_description, candidate_profile)
            
            # Generate resume content using available method (API or local LLM)
            resume_content = self.generate_text(prompt)
            
            # Save to file
            if not output_path:
                # Create a unique filename based on timestamp and job details
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                job_title = candidate_profile.get("target_job_title", "resume")
                output_dir = os.path.join(project_root, "data", "generated_resumes")
                os.makedirs(output_dir, exist_ok=True)
                output_path = os.path.join(output_dir, f"{timestamp}_{job_title.replace(' ', '_')}.docx")
            
            # Save the resume (implementation depends on your format/template needs)
            # This is a simplified example
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(resume_content)
            
            logger.info(f"Generated resume saved to {output_path}")
            return output_path, resume_content
            
        except Exception as e:
            logger.error(f"Error generating resume: {e}")
            # Return a minimal result in case of failure
            return "", "Error generating resume."
            
    def generate_cover_letter(self, 
                       job_description: str, 
                       candidate_resume: str,
                       company_info: str,
                       output_path: Optional[str] = None,
                       template_type: Optional[CoverLetterTemplate] = None,
                       referral_info: Optional[str] = None) -> Tuple[str, str]:
        """
        Generate a personalized cover letter based on job description and resume.
        
        Args:
            job_description: The job description to tailor the cover letter to.
            candidate_resume: The candidate's resume content.
            company_info: Information about the company.
            output_path: Optional path to save the generated cover letter.
                        If None, a default path will be used.
            template_type: Optional template type for the cover letter.
                          If None, a standard template will be used.
            referral_info: Optional referral information to include.
                          
        Returns:
            A tuple containing (file_path, cover_letter_content).
        """
        try:
            prompt = self._prepare_cover_letter_prompt(
                job_description, 
                candidate_resume, 
                company_info,
                template_type or CoverLetterTemplate.STANDARD,
                referral_info
            )
            
            # Generate cover letter content using available method (API or local LLM)
            cover_letter_content = self.generate_text(prompt)
            
            # Save to file
            if not output_path:
                # Create a unique filename based on timestamp and job details
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                job_title = "cover_letter"  # Extract from job description if possible
                if "title" in job_description.lower():
                    job_title = job_description.split("title")[1].split("\n")[0].strip()[:30]
                output_dir = os.path.join(project_root, "data", "generated_cover_letters")
                os.makedirs(output_dir, exist_ok=True)
                output_path = os.path.join(output_dir, f"{timestamp}_{job_title.replace(' ', '_')}.docx")
            
            # Save the cover letter (implementation depends on your format/template needs)
            # This is a simplified example
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(cover_letter_content)
            
            logger.info(f"Generated cover letter saved to {output_path}")
            return output_path, cover_letter_content
            
        except Exception as e:
            logger.error(f"Error generating cover letter: {e}")
            # Return a minimal result in case of failure
            return "", "Error generating cover letter."
    
    def _prepare_resume_prompt(self, 
                        job_description: str, 
                        candidate_profile: Dict[str, Any]) -> str:
        """
        Prepare the prompt for resume generation with enhanced AI analysis.
        """
        # Extract key requirements from job description using AI
        key_requirements = self._extract_job_requirements(job_description)
        
        # Score and sort candidate skills based on job requirements
        scored_skills = self._score_skills_match(candidate_profile.get("skills", []), key_requirements)
        
        # Format candidate profile with prioritized skills
        profile_sections = []
        for section, content in candidate_profile.items():
            if section == "skills":
                # Order skills by relevance score
                skills_text = "\nSkills:\n" + "\n".join(f"- {skill}" for skill, _ in scored_skills)
                profile_sections.append(skills_text)
            else:
                profile_sections.append(f"{section}:\n{content}")
                
        profile_text = "\n\n".join(profile_sections)
        
        # Generate focused prompt
        prompt = self.config.resume_template_prompt.format(
            job_description=job_description,
            candidate_profile=profile_text,
            key_requirements="\n".join(f"- {req}" for req in key_requirements)
        )
        
        return prompt

    def _extract_job_requirements(self, job_description: str) -> List[str]:
        """Extract key requirements from job description using AI analysis."""
        try:
            # Use the LLM to analyze requirements
            analysis_prompt = """
            Analyze the following job description and extract the key requirements including:
            - Technical skills
            - Soft skills
            - Experience level
            - Education requirements
            - Industry knowledge
            
            Job Description:
            {description}
            
            Format each requirement as a concise bullet point.
            """
            
            response = self._generate_text(
                prompt=analysis_prompt.format(description=job_description),
                system_prompt="You are an expert job requirements analyzer."
            )
            
            # Parse bullet points from response
            requirements = [line.strip("- ").strip() 
                          for line in response.split("\n") 
                          if line.strip().startswith("-")]
            
            return requirements
            
        except Exception as e:
            logger.error(f"Error extracting job requirements: {e}")
            return []

    def _score_skills_match(self, 
                      candidate_skills: List[str], 
                      job_requirements: List[str]) -> List[Tuple[str, float]]:
        """Score candidate skills against job requirements using AI matching."""
        try:
            # Prepare scoring prompt
            scoring_prompt = """
            Score how well each candidate skill matches the job requirements.
            Consider direct matches, related skills, and transferable skills.
            
            Job Requirements:
            {requirements}
            
            Candidate Skills:
            {skills}
            
            For each skill, respond with: skill|score
            Score from 0.0 to 1.0 where 1.0 is a perfect match.
            """
            
            formatted_prompt = scoring_prompt.format(
                requirements="\n".join(job_requirements),
                skills="\n".join(candidate_skills)
            )
            
            response = self._generate_text(
                prompt=formatted_prompt,
                system_prompt="You are an expert skills matcher."
            )
            
            # Parse scores and sort
            scored_skills = []
            for line in response.split("\n"):
                if "|" in line:
                    skill, score = line.split("|")
                    try:
                        scored_skills.append((skill.strip(), float(score.strip())))
                    except ValueError:
                        continue
                        
            return sorted(scored_skills, key=lambda x: x[1], reverse=True)
            
        except Exception as e:
            logger.error(f"Error scoring skills match: {e}")
            return [(skill, 0.5) for skill in candidate_skills]
            
    def _prepare_cover_letter_prompt(self,
                              job_description: str,
                              candidate_resume: str,
                              company_info: str,
                              template_type: CoverLetterTemplate,
                              referral_info: Optional[str] = None) -> str:
        """
        Prepare the prompt for cover letter generation with enhanced AI analysis.
        """
        # Use the template prompt from config
        prompt = self.config.cover_letter_template_prompt.format(
            job_description=job_description,
            candidate_resume=candidate_resume,
            company_info=company_info
        )
        
        return prompt
        
    def _generate_text(self, 
                prompt: str, 
                system_prompt: Optional[str] = None, 
                **kwargs) -> str:
        """
        Generate text using the LLM.
        
        Args:
            prompt: The prompt text.
            system_prompt: Optional system prompt for LLM.
            **kwargs: Additional parameters for text generation.
            
        Returns:
            Generated text string.
        """
        if not self.llm:
            logger.error("LLM not initialized")
            return ""
            
        try:
            # Prepare parameters
            params = {
                "temperature": self.config.temperature,
                "top_p": self.config.top_p,
                "max_tokens": self.config.context_length // 2
            }
            params.update(kwargs)
            # Combine system + user prompt if provided
            full_prompt = f"{system_prompt}\n{prompt}" if system_prompt else prompt
            # Call LLM
            response = self.llm(full_prompt, **params)
            # Extract text
            text = ""
            if hasattr(response, 'choices'):
                text = response.choices[0].text
            elif isinstance(response, dict) and 'choices' in response:
                text = response['choices'][0]['text']
            return text.strip()
        except Exception as e:
            logger.error(f"Error generating text with LLM: {e}")
            return ""
            
    def _create_resume_document(self, content: str, output_path: str) -> str:
        """
        Create a resume document from the generated content.
        
        Args:
            content: The generated resume content.
            output_path: Path to save the document.
            
        Returns:
            Path to the created document.
        """
        try:
            # Check if there's a template available
            template_path = "../templates/resume_template.docx"
            if os.path.exists(template_path):
                # Use the template
                return self._create_document_from_template(content, template_path, output_path)
            else:
                # Create a document from scratch
                return self._create_document(content, output_path)
                
        except Exception as e:
            logger.error(f"Error creating resume document: {e}")
            return ""
            
    def _create_cover_letter_document(self, content: str, output_path: str, template_path: Optional[Path] = None) -> str:
        """
        Create a cover letter document from the generated content.
        
        Args:
            content: The generated cover letter content.
            output_path: Path to save the document.
            template_path: Optional specific template path to use.
            
        Returns:
            Path to the created document.
        """
        try:
            # Use provided template path or default
            template_path_str = str(template_path) if template_path else "../templates/cover_letter_template.docx"
            
            if os.path.exists(template_path_str):
                # Use the template
                return self._create_document_from_template(content, template_path_str, output_path)
            else:
                # Create a document from scratch
                return self._create_document(content, output_path)
                
        except Exception as e:
            logger.error(f"Error creating cover letter document: {e}")
            return ""
            
    def _create_document_from_template(self, content: str, template_path: str, output_path: str) -> str:
        """
        Create a document using a template.
        
        Args:
            content: The document content.
            template_path: Path to the template file.
            output_path: Path to save the document.
            
        Returns:
            Path to the created document.
        """
        try:
            # Use DocxTemplate for template-based document creation
            doc = DocxTemplate(template_path)
            
            # Parse the content into sections
            sections = self._parse_content_to_sections(content)
            
            # Add current date to the context
            sections['current_date'] = datetime.now().strftime("%B %d, %Y")
            
            # Render the template with the content
            doc.render(sections)
            
            # Ensure the output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Save the document
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating document from template: {e}")
            return ""
            
    def _create_document(self, content: str, output_path: str) -> str:
        """
        Create a document from scratch.
        
        Args:
            content: The document content.
            output_path: Path to save the document.
            
        Returns:
            Path to the created document.
        """
        try:
            # Create a new document
            doc = Document()
            
            # Set document properties
            doc.styles['Normal'].font.name = 'Calibri'
            doc.styles['Normal'].font.size = Pt(11)
            
            # Split content into paragraphs
            paragraphs = content.split('\n')
            
            # Add paragraphs to the document
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Check if it's a heading (# or ## at the beginning)
                    if paragraph.strip().startswith('#'):
                        level = paragraph.count('#', 0, paragraph.find(' '))
                        heading_text = paragraph.strip('#').strip()
                        
                        if level == 1:
                            doc.add_heading(heading_text, 0)  # Title
                        elif level == 2:
                            doc.add_heading(heading_text, 1)  # Heading 1
                        else:
                            doc.add_heading(heading_text, level - 1)  # Other headings
                    else:
                        # Regular paragraph
                        doc.add_paragraph(paragraph)
            
            # Ensure the output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Save the document
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            return ""
            
    def _parse_content_to_sections(self, content: str) -> Dict[str, Any]:
        """
        Parse the content into sections for template rendering.
        
        Args:
            content: The document content.
            
        Returns:
            A dictionary with sections.
        """
        sections = {
            'content': content  # Default, full content
        }
        
        try:
            # Split content into lines
            lines = content.split('\n')
            
            # Current section and section content
            current_section = None
            section_content = []
            
            # Extract sections based on markdown headings
            for line in lines:
                if line.startswith('# '):
                    # Level 1 heading (title)
                    if current_section:
                        # Save previous section
                        sections[current_section] = '\n'.join(section_content)
                    
                    # Start new section
                    current_section = line[2:].strip().lower().replace(' ', '_')
                    section_content = []
                elif line.startswith('## '):
                    # Level 2 heading (section)
                    if current_section:
                        # Save previous section
                        sections[current_section] = '\n'.join(section_content)
                    
                    # Start new section
                    current_section = line[3:].strip().lower().replace(' ', '_')
                    section_content = []
                else:
                    # Content line
                    if current_section:
                        section_content.append(line)
            
            # Save the last section
            if current_section and section_content:
                sections[current_section] = '\n'.join(section_content)
                
            # Extract specific sections for templates
            # Contact information
            if 'contact_information' in sections:
                contact_lines = sections['contact_information'].split('\n')
                for line in contact_lines:
                    if ':' in line:
                        key, value = line.split(':', 1)
                        key = key.strip().lower().replace(' ', '_')
                        sections[key] = value.strip()
            
            # Skills section - convert to list
            if 'skills' in sections:
                skills_text = sections['skills']
                skills_list = []
                
                for line in skills_text.split('\n'):
                    line = line.strip()
                    if line.startswith('- '):
                        skills_list.append(line[2:])
                    elif line:
                        skills_list.append(line)
                
                sections['skills_list'] = skills_list
                
            # Look for an address block
            address_block = []
            for i, line in enumerate(lines):
                # Look for patterns like a name followed by address details
                if i < 10:  # Usually at the top
                    address_block.append(line)
            
            if address_block:
                sections['address_block'] = '\n'.join(address_block[:5])  # First few lines
            
            # Extract recipient information for cover letters
            recipient_block = []
            for i in range(len(lines)):
                if i < 15 and "hiring manager" in lines[i].lower() or "recruiter" in lines[i].lower():
                    for j in range(i, min(i+5, len(lines))):
                        recipient_block.append(lines[j])
                    break
            
            if recipient_block:
                sections['recipient'] = '\n'.join(recipient_block)
                
        except Exception as e:
            logger.error(f"Error parsing content to sections: {e}")
            
        return sections
    
    def _extract_profile_from_resume(self, resume_content: str) -> Dict[str, Any]:
        """
        Extract a profile dictionary from resume content.
        
        Args:
            resume_content: The resume text content.
            
        Returns:
            A dictionary with profile information.
        """
        profile = {}
        
        try:
            # Simple extraction of main sections
            sections = self._parse_content_to_sections(resume_content)
            
            # Extract name - usually the first heading or first line
            lines = resume_content.split('\n')
            for line in lines[:5]:
                if line.strip() and not line.startswith('#'):
                    profile['name'] = line.strip()
                    break
            
            # Extract other sections
            if 'contact_information' in sections:
                profile['contact'] = sections['contact_information']
            
            if 'skills' in sections:
                profile['skills'] = sections['skills'].split('\n')
            
            if 'experience' in sections:
                profile['experience'] = sections['experience']
            
            if 'education' in sections:
                profile['education'] = sections['education']
                
        except Exception as e:
            logger.error(f"Error extracting profile from resume: {e}")
            
        return profile
    
    def _extract_company_name(self, job_description: str) -> str:
        """
        Extract company name from job description.
        
        Args:
            job_description: Job description text.
            
        Returns:
            Extracted company name or "Unknown".
        """
        try:
            # A very simple heuristic - look for common patterns
            patterns = [
                r"(?i)at\s+([A-Z][A-Za-z0-9\s&,.-]+?)(?=\s+is\s+looking)",
                r"(?i)with\s+([A-Z][A-Za-z0-9\s&,.-]+?)(?=\s+as)",
                r"(?i)join\s+([A-Z][A-Za-z0-9\s&,.-]+?)(?=\s+as)",
                r"(?i)about\s+([A-Z][A-Za-z0-9\s&,.-]+?)(?=\s*:)"
            ]
            
            import re
            for pattern in patterns:
                matches = re.search(pattern, job_description)
                if matches:
                    return matches.group(1).strip()
                    
            # Default if no match found
            return "Unknown"
            
        except Exception as e:
            logger.error(f"Error extracting company name: {e}")
            return "Unknown"
            
    def _extract_job_title(self, job_description: str) -> str:
        """
        Extract job title from job description.
        
        Args:
            job_description: Job description text.
            
        Returns:
            Extracted job title or "Position".
        """
        try:
            # Simple heuristic to find job title
            patterns = [
                r"(?i)(Software Engineer|Data Scientist|Machine Learning Engineer|Full Stack Developer|Frontend Developer|Backend Developer|DevOps Engineer|Cloud Engineer|AI Engineer|ML Engineer)(?=\s+at)",
                r"(?i)(Software Engineer|Data Scientist|Machine Learning Engineer|Full Stack Developer|Frontend Developer|Backend Developer|DevOps Engineer|Cloud Engineer|AI Engineer|ML Engineer)(?=\s+position)"
            ]
            
            import re
            for pattern in patterns:
                matches = re.search(pattern, job_description)
                if matches:
                    return matches.group(1).strip()
                    
            # Look at the beginning of the job description
            first_line = job_description.split('\n')[0]
            if len(first_line) < 100:  # Usually the title is short
                return first_line.strip()
                
            # Default if no match found
            return "Position"
            
        except Exception as e:
            logger.error(f"Error extracting job title: {e}")
            return "Position"


# Example usage
def main():
    # Example job description
    job_description = """
    Software Engineer at TechCorp
    
    We're looking for a skilled Software Engineer to join our team at TechCorp. 
    The ideal candidate will have experience with Python, JavaScript, and cloud technologies.
    
    Responsibilities:
    - Develop and maintain web applications
    - Write clean, efficient, and maintainable code
    - Collaborate with cross-functional teams
    
    Requirements:
    - Bachelor's degree in Computer Science or related field
    - 2+ years of experience in software development
    - Proficiency in Python and JavaScript
    - Experience with cloud platforms (AWS, Azure, or GCP)
    - Strong problem-solving skills
    """
    
    # Example candidate profile
    candidate_profile = {
        "name": "Jane Doe",
        "email": "jane.doe@email.com",
        "phone": "555-123-4567",
        "summary": "Software Engineer with 3 years of experience in Python and JavaScript development.",
        "skills": ["Python", "JavaScript", "React", "Node.js", "AWS", "Docker"],
        "experience": [
            {
                "title": "Software Engineer",
                "company": "Tech Solutions Inc.",
                "duration": "2021-2023",
                "description": "Developed and maintained web applications using Python and JavaScript."
            },
            {
                "title": "Junior Developer",
                "company": "StartUp Co.",
                "duration": "2020-2021",
                "description": "Assisted in building frontend components with React."
            }
        ],
        "education": [
            {
                "degree": "B.S. Computer Science",
                "institution": "State University",
                "year": "2020"
            }
        ]
    }
    
    # Example company info
    company_info = """
    TechCorp is a leading technology company specializing in cloud solutions and web applications.
    Founded in 2010, the company has grown to over 500 employees and serves clients worldwide.
    The company culture emphasizes innovation, collaboration, and work-life balance.
    """
    
    # Create resume generator
    resume_generator = ResumeGenerator()
    
    # Generate resume
    resume_path, resume_content = resume_generator.generate_resume(
        job_description=job_description,
        candidate_profile=candidate_profile
    )
    
    print(f"Resume saved to: {resume_path}")
    
    # Generate standard cover letter
    cover_letter_path, cover_letter_content = resume_generator.generate_cover_letter(
        job_description=job_description,
        candidate_resume=resume_content,
        company_info=company_info,
        template_type=CoverLetterTemplate.STANDARD
    )
    
    print(f"Cover letter saved to: {cover_letter_path}")
    
    # Generate technical cover letter
    tech_cover_letter_path, tech_cover_letter_content = resume_generator.generate_cover_letter(
        job_description=job_description,
        candidate_resume=resume_content,
        company_info=company_info,
        template_type=CoverLetterTemplate.TECHNICAL
    )
    
    print(f"Technical cover letter saved to: {tech_cover_letter_path}")
    
    # Generate referral cover letter with referral info
    referral_info = "John Smith, Senior Engineer at TechCorp, who I worked with at my previous company."
    referral_cover_letter_path, referral_cover_letter_content = resume_generator.generate_cover_letter(
        job_description=job_description,
        candidate_resume=resume_content,
        company_info=company_info,
        template_type=CoverLetterTemplate.REFERRAL,
        referral_info=referral_info
    )
    
    print(f"Referral cover letter saved to: {referral_cover_letter_path}")


if __name__ == "__main__":
    main()